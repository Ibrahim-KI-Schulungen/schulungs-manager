#!/usr/bin/env python3
"""
Vertrag-Generator für Schulungs-Extraktor
Generiert Beauftragungsverträge und Rahmenverträge aus Vorlagen.

Unterstützte Vertragstypen:
1. Beauftragungsvertrag - für einzelne Schulungsaufträge
2. Rahmenvertrag - für neue Trainer-Kooperationen
"""

import os
import json
import re
from datetime import datetime
from typing import Dict, Any, Optional, List
from docx import Document


class VertragGenerator:
    """Generiert Beauftragungsverträge aus Vorlage."""

    def __init__(self, config_path: str = None, vorlagen_dir: str = None, vertraege_dir: str = None):
        self.script_dir = os.path.dirname(os.path.abspath(__file__))
        self.vorlagen_dir = vorlagen_dir or os.path.join(self.script_dir, "vorlagen")
        self.vertraege_dir = vertraege_dir or os.path.join(self.script_dir, "vertraege")
        self.vorlage_pfad = os.path.join(self.vorlagen_dir, "Beauftragungsvertrag_Vorlage.docx")
        self.rahmenvertrag_vorlage_pfad = os.path.join(self.vorlagen_dir, "Rahmenvertrag_Vorlage.docx")
        self.config = self._lade_config(config_path)
        self.warnungen = []

    def get_trainer_liste(self) -> List[Dict]:
        """Gibt Liste aller bekannten Trainer zurück."""
        return self.config.get("bekannte_trainer", [])

    def _lade_config(self, config_path: str = None) -> Dict:
        """Lädt die Konfigurationsdatei."""
        if config_path is None:
            config_path = os.path.join(self.script_dir, "config.json")

        try:
            with open(config_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except FileNotFoundError:
            return {"bekannte_trainer": []}

    def _finde_trainer(self, trainer_name: str) -> Optional[Dict]:
        """Findet Trainer in config anhand Name oder Kurzname."""
        if not trainer_name:
            return None

        trainer_lower = trainer_name.lower().strip()

        for trainer in self.config.get("bekannte_trainer", []):
            # Prüfe vollständigen Namen
            if trainer.get("name", "").lower() == trainer_lower:
                return trainer
            # Prüfe Kurznamen
            for kurzname in trainer.get("kurznamen", []):
                if kurzname.lower() == trainer_lower:
                    return trainer

        return None

    def _datum_zu_text(self, datum_str: str) -> str:
        """Konvertiert DD.MM.YYYY zu 'TT. Monat JJJJ'."""
        if not datum_str:
            return "[FEHLT]"

        monate = {
            1: "Januar", 2: "Februar", 3: "März", 4: "April",
            5: "Mai", 6: "Juni", 7: "Juli", 8: "August",
            9: "September", 10: "Oktober", 11: "November", 12: "Dezember"
        }

        try:
            dt = datetime.strptime(datum_str, "%d.%m.%Y")
            return f"{dt.day}. {monate[dt.month]} {dt.year}"
        except ValueError:
            return datum_str

    def _berechne_schulungstage(self, datum_start: str, datum_ende: str) -> int:
        """Berechnet Anzahl der Schulungstage."""
        if not datum_start or not datum_ende:
            return 1

        try:
            start = datetime.strptime(datum_start, "%d.%m.%Y")
            ende = datetime.strptime(datum_ende, "%d.%m.%Y")
            tage = (ende - start).days + 1
            return max(1, tage)
        except ValueError:
            return 1

    def _trainer_nachname(self, trainer_name: str) -> str:
        """Extrahiert Nachnamen aus Trainer-Name."""
        if not trainer_name:
            return "Unbekannt"

        # Bei UG/GmbH den ersten Teil nehmen
        if "UG" in trainer_name or "GmbH" in trainer_name:
            parts = trainer_name.split()
            return parts[0] if parts else "Unbekannt"

        # Letztes Wort ist Nachname
        teile = trainer_name.split()
        return teile[-1] if teile else "Unbekannt"

    def _reisekosten_text(self, reisekosten: Any) -> str:
        """Konvertiert Reisekosten zu Vertragstext."""
        if reisekosten is None:
            return "[REISEKOSTEN PRÜFEN]"
        if isinstance(reisekosten, str):
            reisekosten_lower = reisekosten.lower()
            if "inkl" in reisekosten_lower:
                return "Im Honorar enthalten"
            if "erstattet" in reisekosten_lower or "fahrtkosten" in reisekosten_lower or "hotel" in reisekosten_lower:
                return "Fahrtkosten + Hotel werden erstattet"
            if "keine" in reisekosten_lower or reisekosten == "0":
                return "Keine Reisekosten"
            return reisekosten
        if isinstance(reisekosten, (int, float)):
            if reisekosten == 0:
                return "Keine Reisekosten"
            return f"{int(reisekosten)}€ Reisekosten"
        return str(reisekosten)

    def _ersetze_in_paragraph(self, paragraph, ersetzungen: Dict[str, str]):
        """Ersetzt Text in einem Paragraph unter Beibehaltung der Formatierung."""
        if not paragraph.runs:
            return

        # Gesamten Text des Paragraphen zusammensetzen
        full_text = paragraph.text

        # Prüfen ob eine Ersetzung nötig ist
        needs_replacement = False
        for alt, neu in ersetzungen.items():
            if alt in full_text:
                needs_replacement = True
                break

        if not needs_replacement:
            return

        # Ersetzungen durchführen
        for alt, neu in ersetzungen.items():
            full_text = full_text.replace(alt, neu)

        # Text in Runs aktualisieren
        # Strategie: Ersten Run mit neuem Text füllen, Rest leeren
        if paragraph.runs:
            # Formatierung des ersten Runs merken
            first_run = paragraph.runs[0]
            first_run.text = full_text

            # Restliche Runs leeren
            for run in paragraph.runs[1:]:
                run.text = ""

    def _ersetze_in_dokument(self, doc: Document, ersetzungen: Dict[str, str]):
        """Ersetzt Text im gesamten Dokument."""
        # Paragraphen durchgehen
        for paragraph in doc.paragraphs:
            self._ersetze_in_paragraph(paragraph, ersetzungen)

        # Tabellen durchgehen
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._ersetze_in_paragraph(paragraph, ersetzungen)

    def generiere(self, daten: Dict[str, Any]) -> Dict[str, Any]:
        """Generiert Vertrag aus extrahierten Daten."""
        self.warnungen = []

        # Prüfe ob Vorlage existiert
        if not os.path.exists(self.vorlage_pfad):
            return {
                "erfolg": False,
                "fehler": f"Vorlage nicht gefunden unter: {self.vorlage_pfad}\n"
                         f"Bitte Beauftragungsvertrag_AI-Z_Group_final.docx in vorlagen/ ablegen."
            }

        # Trainer-Daten holen
        trainer_name = daten.get("trainer", "")
        trainer_daten = self._finde_trainer(trainer_name)

        if trainer_daten:
            trainer_vollname = trainer_daten.get("name", "[TRAINER-NAME]")
            trainer_strasse = trainer_daten.get("strasse", "[TRAINER-STRASSE]") + ","
            trainer_plz_ort = f"{trainer_daten.get('plz', '[PLZ]')} {trainer_daten.get('ort', '[ORT]')}"
        else:
            self.warnungen.append(
                f"Trainer '{trainer_name}' nicht in config.json gefunden.\n"
                f"  Vertrag wird mit Platzhaltern erstellt. Bitte manuell ergänzen."
            )
            trainer_vollname = "[TRAINER-NAME]"
            trainer_strasse = "[TRAINER-STRASSE],"
            trainer_plz_ort = "[TRAINER-PLZ-ORT]"

        # Extrahierte Daten vorbereiten
        schulungsname = daten.get("schulungsname", "[FEHLT]")
        if not schulungsname:
            schulungsname = "[FEHLT]"
            self.warnungen.append("Feld 'schulungsname' fehlt in Extraktion")

        kunde = daten.get("kunde") or daten.get("firma_ort", "[FEHLT]")
        if not kunde:
            kunde = "[FEHLT]"
            self.warnungen.append("Feld 'kunde' oder 'firma_ort' fehlt in Extraktion")

        format_text = daten.get("format", "vor Ort")

        datum_start = daten.get("datum_start", "")
        datum_ende = daten.get("datum_ende", "")
        datum_start_text = self._datum_zu_text(datum_start)
        datum_ende_text = self._datum_zu_text(datum_ende)
        datum_bereich = f"{datum_start_text} → {datum_ende_text}"

        schulungstage = self._berechne_schulungstage(datum_start, datum_ende)
        schulungstage_text = f"{schulungstage} Schulungstag" if schulungstage == 1 else f"{schulungstage} Schulungstage"

        trainer_kosten = daten.get("trainer_kosten")
        if trainer_kosten:
            trainer_kosten_text = f"{int(trainer_kosten)}€"
        else:
            trainer_kosten_text = "[FEHLT]€"
            self.warnungen.append("Feld 'trainer_kosten' fehlt in Extraktion")

        vorbereitungspauschale = daten.get("vorbereitungspauschale")
        if vorbereitungspauschale:
            vorbereitung_text = f"{int(vorbereitungspauschale)}€ Vorbereitung"
        else:
            vorbereitung_text = "[VORBEREITUNG PRÜFEN]"

        reisekosten = daten.get("reisekosten")
        reisekosten_text = self._reisekosten_text(reisekosten)

        # Ersetzungen definieren
        ersetzungen = {
            "Sebastian Grell": trainer_vollname,
            "Dernburgstr. 9,": trainer_strasse,
            "14057 Berlin": trainer_plz_ort,
            "Einführung in KI-gestützte Marktforschung": schulungsname,
            "HAIX Schuhe Produktions & Vertriebs GmbH": kunde,
            "21. Januar 2026 → 23. Januar 2026": datum_bereich,
            "3 Schulungstag": schulungstage_text,
            "550€": trainer_kosten_text,
            "600€ Vorbereitung": vorbereitung_text,
            "Fahrtkosten + Hotel werden erstattet": reisekosten_text,
        }

        # Format nur ersetzen wenn es "vor Ort" im Dokument gibt
        # (nicht das gesamte Wort "vor Ort" überall ersetzen)

        try:
            # Dokument laden
            doc = Document(self.vorlage_pfad)

            # Ersetzungen durchführen
            self._ersetze_in_dokument(doc, ersetzungen)

            # Dateiname generieren
            trainer_nachname = self._trainer_nachname(trainer_vollname)
            datum_kurz = datum_start if datum_start else "DATUM"
            dateiname = f"Beauftragung_{trainer_nachname}_{datum_kurz}.docx"
            ausgabe_pfad = os.path.join(self.vertraege_dir, dateiname)

            # Speichern
            doc.save(ausgabe_pfad)

            return {
                "erfolg": True,
                "pfad": ausgabe_pfad,
                "dateiname": dateiname,
                "warnungen": self.warnungen,
            }

        except Exception as e:
            return {
                "erfolg": False,
                "fehler": f"Fehler beim Generieren: {str(e)}"
            }

    def generiere_rahmenvertrag(self, trainer_name: str) -> Dict[str, Any]:
        """
        Generiert Rahmenvertrag für einen Trainer.

        Args:
            trainer_name: Name oder Kurzname des Trainers aus config.json

        Returns:
            Dict mit erfolg, pfad, dateiname, warnungen oder fehler
        """
        self.warnungen = []

        # Prüfe ob Vorlage existiert
        if not os.path.exists(self.rahmenvertrag_vorlage_pfad):
            return {
                "erfolg": False,
                "fehler": f"Rahmenvertrag-Vorlage nicht gefunden unter: {self.rahmenvertrag_vorlage_pfad}"
            }

        # Trainer-Daten holen
        trainer_daten = self._finde_trainer(trainer_name)

        if not trainer_daten:
            return {
                "erfolg": False,
                "fehler": f"Trainer '{trainer_name}' nicht in config.json gefunden.\n"
                         f"Bitte zuerst in config.json hinzufügen mit Name, Straße, PLZ und Ort."
            }

        # Trainer-Daten extrahieren
        trainer_vollname = trainer_daten.get("name", "[NAME]")
        trainer_strasse = trainer_daten.get("strasse", "[STRASSE]")
        trainer_plz = trainer_daten.get("plz", "[PLZ]")
        trainer_ort = trainer_daten.get("ort", "[ORT]")
        trainer_plz_ort = f"{trainer_plz} {trainer_ort}"

        # Prüfe ob alle Adressdaten vorhanden sind
        if not trainer_daten.get("strasse"):
            self.warnungen.append(f"Straße fehlt für {trainer_vollname} in config.json")
        if not trainer_daten.get("plz"):
            self.warnungen.append(f"PLZ fehlt für {trainer_vollname} in config.json")
        if not trainer_daten.get("ort"):
            self.warnungen.append(f"Ort fehlt für {trainer_vollname} in config.json")

        # Ersetzungen definieren
        ersetzungen = {
            "{{NAME}}": trainer_vollname,
            "{{STRASSE}}": trainer_strasse,
            "{{PLZ_ORT}}": trainer_plz_ort,
        }

        try:
            # Dokument laden
            doc = Document(self.rahmenvertrag_vorlage_pfad)

            # Ersetzungen durchführen
            self._ersetze_in_dokument(doc, ersetzungen)

            # Vertraege-Ordner erstellen falls nicht vorhanden
            os.makedirs(self.vertraege_dir, exist_ok=True)

            # Dateiname generieren
            trainer_nachname = self._trainer_nachname(trainer_vollname)
            datum_heute = datetime.now().strftime("%d.%m.%Y")
            dateiname = f"Rahmenvertrag_{trainer_nachname}_{datum_heute}.docx"
            ausgabe_pfad = os.path.join(self.vertraege_dir, dateiname)

            # Speichern
            doc.save(ausgabe_pfad)

            return {
                "erfolg": True,
                "pfad": ausgabe_pfad,
                "dateiname": dateiname,
                "trainer": trainer_vollname,
                "warnungen": self.warnungen,
            }

        except Exception as e:
            return {
                "erfolg": False,
                "fehler": f"Fehler beim Generieren des Rahmenvertrags: {str(e)}"
            }


def interaktive_cli():
    """Interaktive CLI für Vertragsgenerierung."""
    print("=" * 60)
    print("VERTRAG-GENERATOR")
    print("=" * 60)
    print()
    print("Welchen Vertrag möchtest du erstellen?")
    print()
    print("  [1] Beauftragungsvertrag (für einzelne Schulung)")
    print("  [2] Rahmenvertrag (für neue Trainer-Kooperation)")
    print()

    while True:
        wahl = input("Deine Wahl (1/2): ").strip()
        if wahl in ["1", "2"]:
            break
        print("Bitte 1 oder 2 eingeben.")

    generator = VertragGenerator()

    if wahl == "1":
        # Beauftragungsvertrag - benötigt Schulungsdaten
        print()
        print("-" * 60)
        print("BEAUFTRAGUNGSVERTRAG")
        print("-" * 60)
        print()
        print("Für Beauftragungsverträge werden Schulungsdaten benötigt.")
        print("Nutze den Schulungs-Extraktor oder gib die Daten manuell ein.")
        print()
        print("Beispiel-Aufruf aus Python:")
        print("  generator = VertragGenerator()")
        print("  ergebnis = generator.generiere(schulungsdaten)")
        print()
        print("Oder nutze Claude Code: 'Erstell einen Beauftragungsvertrag für...'")

    elif wahl == "2":
        # Rahmenvertrag
        print()
        print("-" * 60)
        print("RAHMENVERTRAG")
        print("-" * 60)
        print()

        # Trainer-Liste anzeigen
        trainer_liste = generator.get_trainer_liste()

        if not trainer_liste:
            print("Keine Trainer in config.json gefunden!")
            return

        print("Verfügbare Trainer:")
        print()
        for i, trainer in enumerate(trainer_liste, 1):
            name = trainer.get("name", "Unbekannt")
            ort = trainer.get("ort", "?")
            print(f"  [{i}] {name} ({ort})")
        print()

        # Trainer auswählen
        while True:
            eingabe = input(f"Trainer wählen (1-{len(trainer_liste)}): ").strip()
            try:
                idx = int(eingabe) - 1
                if 0 <= idx < len(trainer_liste):
                    break
            except ValueError:
                pass
            print(f"Bitte eine Zahl zwischen 1 und {len(trainer_liste)} eingeben.")

        ausgewaehlter_trainer = trainer_liste[idx]
        trainer_name = ausgewaehlter_trainer.get("name")

        print()
        print(f"Generiere Rahmenvertrag für: {trainer_name}")
        print()

        # Vertrag generieren
        ergebnis = generator.generiere_rahmenvertrag(trainer_name)

        if ergebnis["erfolg"]:
            print("=" * 60)
            print("✓ RAHMENVERTRAG ERSTELLT")
            print("=" * 60)
            print()
            print(f"  Trainer:  {ergebnis['trainer']}")
            print(f"  Datei:    {ergebnis['dateiname']}")
            print(f"  Pfad:     {ergebnis['pfad']}")

            if ergebnis.get("warnungen"):
                print()
                print("  Warnungen:")
                for w in ergebnis["warnungen"]:
                    print(f"    ⚠ {w}")

            print()
            print("Vertrag öffnen? (j/n): ", end="")
            if input().strip().lower() in ["j", "ja", "y", "yes"]:
                import subprocess
                subprocess.run(["open", ergebnis["pfad"]])

        else:
            print(f"✗ Fehler: {ergebnis['fehler']}")


def teste_generator():
    """Testet den Vertrag-Generator (für Entwicklung)."""
    print("=" * 60)
    print("VERTRAG-GENERATOR - Test")
    print("=" * 60)

    # Test-Daten
    test_daten = {
        "schulungsname": "Einführung in KI-gestützte Marktforschung",
        "datum_start": "09.12.2025",
        "datum_ende": "11.12.2025",
        "trainer": "Lukas Sontheimer",
        "trainer_kosten": 550.0,
        "vorbereitungspauschale": None,
        "reisekosten": None,
        "kunde": "HAIX Schuhe Produktions & Vertriebs GmbH",
        "format": "Vor Ort",
    }

    generator = VertragGenerator()

    print("\nTest-Daten:")
    for k, v in test_daten.items():
        print(f"  {k}: {v}")

    print("\nGeneriere Vertrag...")
    ergebnis = generator.generiere(test_daten)

    if ergebnis["erfolg"]:
        print(f"\n✓ Vertrag erstellt: {ergebnis['pfad']}")
        if ergebnis.get("warnungen"):
            print("\nWarnungen:")
            for w in ergebnis["warnungen"]:
                print(f"  ⚠ {w}")
    else:
        print(f"\n✗ Fehler: {ergebnis['fehler']}")


if __name__ == "__main__":
    import sys
    if len(sys.argv) > 1 and sys.argv[1] == "--test":
        teste_generator()
    else:
        interaktive_cli()
